from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Corporate_Wallet_Digital_Twin_V3_1_System_Dossier.md"
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "assets"
OUTPUT = OUT_DIR / "Corporate_Wallet_Digital_Twin_V3_1_System_Dossier.docx"

NAVY = "0B1F3A"
BLUE = "1769AA"
SKY = "DCEBF7"
GOLD = "D8A848"
INK = "263238"
MUTED = "5D6975"
PALE = "F2F4F7"
TABLE_HEADER = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"
RED = "A33A3A"
GREEN = "287A55"


def pil_color(value: str) -> str:
    return value if value.startswith("#") else f"#{value}"


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    if bold:
        candidates = [
            Path("C:/Windows/Fonts/seguisb.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
        ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded_box(draw: ImageDraw.ImageDraw, xy, fill, outline=BLUE, radius=22, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=pil_color(fill), outline=pil_color(outline), width=width)


def centered_lines(draw: ImageDraw.ImageDraw, box, lines: Iterable[str], title_color=NAVY, body_color=INK):
    x1, y1, x2, y2 = box
    lines = list(lines)
    title_f = font(30, True)
    body_f = font(21)
    heights = [38] + [30] * (len(lines) - 1)
    total = sum(heights)
    y = y1 + (y2 - y1 - total) / 2
    for idx, line in enumerate(lines):
        f = title_f if idx == 0 else body_f
        color = title_color if idx == 0 else body_color
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text(((x1 + x2 - (bbox[2] - bbox[0])) / 2, y), line, font=f, fill=pil_color(color))
        y += heights[idx]


def arrow(draw: ImageDraw.ImageDraw, start, end, color=BLUE, width=5):
    draw.line([start, end], fill=pil_color(color), width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    back = 18
    wing = 9
    p1 = (x2 - back * ux + wing * px, y2 - back * uy + wing * py)
    p2 = (x2 - back * ux - wing * px, y2 - back * uy - wing * py)
    draw.polygon([end, p1, p2], fill=pil_color(color))


def build_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 1040), pil_color(WHITE))
    draw = ImageDraw.Draw(img)
    draw.text((80, 48), "Corporate Wallet Digital Twin V3 â€” target production flow", font=font(40, True), fill=pil_color(NAVY))
    draw.text((80, 102), "Private, point-in-time and governed from source to outcome", font=font(24), fill=pil_color(MUTED))

    boxes = {
        "sources": (70, 210, 360, 780),
        "ingest": (440, 210, 735, 455),
        "lake": (440, 535, 735, 780),
        "models": (820, 210, 1140, 455),
        "evidence": (820, 535, 1140, 780),
        "bff": (1225, 210, 1535, 455),
        "events": (1225, 535, 1535, 780),
        "controls": (1615, 210, 1850, 780),
    }
    fills = {"sources": PALE, "ingest": SKY, "lake": SKY, "models": "E7F4EE", "evidence": "FFF5DA", "bff": "E7F4EE", "events": "F3EAF8", "controls": "FCEAEA"}
    for name, box in boxes.items():
        rounded_box(draw, box, fills[name], NAVY if name in {"sources", "controls"} else BLUE)

    centered_lines(draw, boxes["sources"], ["Sources", "Bank activity & balances", "CRM and finance", "Public documents", "Consented multibank data"])
    centered_lines(draw, boxes["ingest"], ["AWS ingestion", "Contract validation", "Quarantine", "Immutable source hashes"])
    centered_lines(draw, boxes["lake"], ["S3 + Delta Lake", "Raw â†’ conformed â†’ curated", "Point-in-time features", "Object-locked snapshots"])
    centered_lines(draw, boxes["models"], ["Decision services", "Bounds + posterior + shadow", "PU + temporal dynamics", "CVaR portfolio + VOI"])
    centered_lines(draw, boxes["evidence"], ["Evidence + GenAI", "Extraction candidates", "Four-eyes approval", "Controlled narration"])
    centered_lines(draw, boxes["bff"], ["Entitled workbench", "ABAC server-side APIs", "Separate claim layers", "CRM adapter"])
    centered_lines(draw, boxes["events"], ["MSK learning loop", "Eligibility and assignment", "Interaction and action", "Outcome and access events"])
    centered_lines(draw, boxes["controls"], ["Control plane", "SSO + OPA + Unity Catalog", "KMS + MLflow registry", "OpenTelemetry + SIEM", "Monitoring + rollback"])

    arrow(draw, (360, 335), (440, 335))
    arrow(draw, (587, 455), (587, 535))
    arrow(draw, (735, 335), (820, 335))
    arrow(draw, (735, 655), (820, 655))
    arrow(draw, (1140, 335), (1225, 335))
    arrow(draw, (1380, 455), (1380, 535))
    arrow(draw, (1535, 655), (1615, 655))
    arrow(draw, (1615, 335), (1535, 335), color=GOLD)
    draw.text((70, 900), "Release truth", font=font(25, True), fill=pil_color(NAVY))
    draw.text((235, 900), "Client demo READY", font=font(25, True), fill=pil_color(GREEN))
    draw.text((515, 900), "Bank production NOT_PROMOTABLE until authority-dependent gates pass", font=font(25, True), fill=pil_color(RED))
    img.save(path, quality=95)


def build_ladder_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 900), pil_color(WHITE))
    draw = ImageDraw.Draw(img)
    draw.text((80, 45), "Evidence and claim ladder", font=font(40, True), fill=pil_color(NAVY))
    draw.text((80, 98), "Stronger evidence permits stronger labels â€” never automatic promotion", font=font(24), fill=pil_color(MUTED))
    tiers = [
        ("E0", "Governed prior", "Prior-led posterior or scenario"),
        ("E1", "Audited public evidence", "Noisy/censored anchor"),
        ("E2", "Client / RM attestation", "Client-validated inference"),
        ("E3", "Multibank observation", "Measured share eligible"),
        ("E4", "Reconciled economics / outcome", "Reconciled or causal evidence eligible"),
    ]
    colors = ["EEF2F5", "DFECF7", "D7EBDD", "F8E9C5", "F1D6D6"]
    left = 80
    base_y = 600
    step_w = 330
    for i, (tier, name, claim) in enumerate(tiers):
        x1 = left + i * 345
        y1 = base_y - i * 100
        x2 = x1 + step_w
        y2 = base_y + 180
        rounded_box(draw, (x1, y1, x2, y2), colors[i], NAVY, radius=18, width=3)
        draw.text((x1 + 20, y1 + 18), tier, font=font(34, True), fill=pil_color(NAVY))
        draw.text((x1 + 20, y1 + 65), name, font=font(22, True), fill=pil_color(BLUE))
        wrapped = wrap_text(draw, claim, font(19), step_w - 40)
        draw.multiline_text((x1 + 20, y1 + 105), wrapped, font=font(19), fill=pil_color(INK), spacing=6)
    draw.text((82, 828), "Observed â†’ identified bound â†’ posterior â†’ scenario â†’ causal", font=font(23, True), fill=pil_color(NAVY))
    img.save(path, quality=95)


def build_decision_architecture(path: Path) -> None:
    """Render the additive governed-substrate to V3-decision flow."""
    img = Image.new("RGB", (1900, 1120), pil_color(WHITE))
    draw = ImageDraw.Draw(img)
    draw.text((80, 42), "Corporate Wallet Digital Twin V3 - governed decision loop", font=font(40, True), fill=pil_color(NAVY))
    draw.text((80, 96), "Evidence and claim class are preserved while uncertainty moves into decisions", font=font(24), fill=pil_color(MUTED))

    items = [
        ("Observed + evidence", "Syn Bank activity", "E1/E2/E3 facts"),
        ("Bounds + posterior", "Independent support", "Product distributions"),
        ("Shadow Wallet", "Anonymous flows", "Exact mass balance"),
        ("Need + dynamics", "PU need", "Change / leakage"),
        ("Robust portfolio", "Mean + lower CVaR", "Capacity constraints"),
        ("Decision VOI", "Cost + latency", "Approval-gated queue"),
        ("Cited brief", "Sealed claim pack", "Deterministic fallback"),
    ]
    fills = [PALE, SKY, "E7F4EE", "FFF5DA", "F3EAF8", "FCEAEA", "E7F4EE"]
    xs = [55, 315, 575, 835, 1095, 1355, 1615]
    for idx, ((title, line1, line2), fill) in enumerate(zip(items, fills)):
        box = (xs[idx], 270, xs[idx] + 230, 600)
        rounded_box(draw, box, fill, NAVY, radius=18, width=3)
        centered_lines(draw, box, [title, line1, line2])
        if idx < len(items) - 1:
            arrow(draw, (xs[idx] + 230, 435), (xs[idx] + 260, 435), color=BLUE, width=4)

    rounded_box(draw, (160, 735, 1740, 965), PALE, BLUE, radius=20, width=3)
    centered_lines(
        draw,
        (160, 735, 1740, 965),
        [
            "Governance envelope",
            "Point-in-time + artifact lineage + deny-by-default ABAC",
            "No named competitor | No measured share | No confirmed leakage | No causal value",
            "Eligibility, assignment, interaction, action and outcome events close the learning loop",
        ],
    )
    arrow(draw, (1735, 600), (1735, 735), color=GOLD)
    arrow(draw, (160, 850), (90, 850), color=GOLD)
    arrow(draw, (90, 850), (90, 600), color=GOLD)
    draw.text((80, 1035), "Release: V3 Decision Lab READY | Bank production NOT_PROMOTABLE", font=font(25, True), fill=pil_color(RED))
    img.save(path, quality=95)


def wrap_text(draw, text, fnt, max_width):
    words = text.split()
    lines = []
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        if draw.textbbox((0, 0), trial, font=fnt)[2] <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    return "\n".join(lines)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_alt_text(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    # A valid field result prevents placeholder prose from surfacing in viewers
    # that do not refresh fields on open. Word/LibreOffice update the value.
    text.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2, text, fld_char3])


def add_numbering_definition(doc: Document):
    """Install compact_reference_guide level-0 bullet and decimal definitions."""
    numbering = doc.part.numbering_part.element
    max_abs = max([int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))] or [0])
    max_num = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] or [0])

    ids = {}
    for kind, fmt, text_value in (("bullet", "bullet", "â€¢"), ("decimal", "decimal", "%1.")):
        max_abs += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(max_abs))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        if kind == "bullet":
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), "Calibri")
            fonts.set(qn("w:hAnsi"), "Calibri")
            r_pr.append(fonts)
            lvl.append(r_pr)
        abstract.append(lvl)
        numbering.append(abstract)

        max_num += 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(max_num))
        abs_id = OxmlElement("w:abstractNumId")
        abs_id.set(qn("w:val"), str(max_abs))
        num.append(abs_id)
        numbering.append(num)
        ids[kind] = max_num
    return ids


def new_number_instance(doc: Document, template_num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    template = None
    for num in numbering.findall(qn("w:num")):
        if int(num.get(qn("w:numId"))) == template_num_id:
            template = num
            break
    if template is None:
        return template_num_id
    abstract = template.find(qn("w:abstractNumId"))
    abstract_id = abstract.get(qn("w:val"))
    next_id = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), abstract_id)
    num.append(abs_id)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return next_id


def set_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, numid])


def set_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def set_page_break_before(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(OxmlElement("w:pageBreakBefore"))


def add_inline(paragraph, text: str, *, color=None, size=None, font_name=None):
    token_re = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|https?://\S+)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            r = paragraph.add_run(text[pos:match.start()])
            style_run(r, color, size, font_name)
        token = match.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(size or 9.5)
            r.font.color.rgb = RGBColor.from_string(NAVY)
        elif token.startswith("["):
            label = token[1:token.index("]")]
            url = token[token.index("(") + 1:-1]
            add_hyperlink(paragraph, label, url)
            r = None
        else:
            add_hyperlink(paragraph, token.rstrip(".,"), token.rstrip(".,"))
            punctuation = token[len(token.rstrip(".,")):]
            if punctuation:
                paragraph.add_run(punctuation)
            r = None
        if r is not None:
            style_run(r, color, size, font_name)
        pos = match.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        style_run(r, color, size, font_name)


def style_run(run, color=None, size=None, font_name=None):
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(doc: Document):
    """Apply the exact compact_reference_guide preset tokens."""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string("1F4D78")
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    if "Caption" in styles:
        cap = styles["Caption"]
    else:
        cap = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    cap.font.name = "Calibri"
    cap.font.size = Pt(8.5)
    cap.font.italic = True
    cap.font.color.rgb = RGBColor.from_string(MUTED)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)

    code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH) if "Code Block" not in styles else styles["Code Block"]
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0


def configure_section(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    return section


def add_running_header_footer(section):
    def populate_header(header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("CORPORATE WALLET DIGITAL TWIN V3  /  SYSTEM DOSSIER")
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(MUTED)

    def populate_footer(footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("IMPLEMENTATION SNAPSHOT  â€¢  10 AUGUST 2026     ")
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        add_field(p, "PAGE")

    populate_header(section.header)
    populate_header(section.even_page_header)
    populate_footer(section.footer)
    populate_footer(section.even_page_footer)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(50)
    r = p.add_run("STANDARD BANK HACKATHON 2026  /  V3.1 SYSTEM RECORD")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Corporate Wallet\nDigital Twin V3.1")
    r.font.name = "Calibri"
    r.font.size = Pt(32)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(32)
    r = p.add_run("Complete System Dossier")
    r.font.name = "Calibri"
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.7)
    p.paragraph_format.right_indent = Inches(0.7)
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run("The authoritative record of the Business Model Twin, governed conversations, wallet substrate, dual value, decision-focused learning, controlled GenAI, production engineering, validation and bank handoff.")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(INK)

    table = doc.add_table(rows=6, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [
        ("DOCUMENT CONTROL", "VALUE"),
        ("Team", "Corporate Wallet Digital Twin"),
        ("Team member", "Christopher Koen"),
        ("Version", "3.1.0"),
        ("As of", "12 August 2026"),
        ("Release truth", "V3.1 Decision Twin READY / Bank production NOT_PROMOTABLE"),
    ]
    set_table_width(table, [2200, 7160])
    repeat_table_header(table.rows[0])
    for row_idx, (row, (label, value)) in enumerate(zip(table.rows, values)):
        set_cell_shading(row.cells[0], NAVY)
        set_cell_shading(row.cells[1], NAVY if row_idx == 0 else PALE)
        set_cell_margins(row.cells[0], 130, 120, 130, 120)
        set_cell_margins(row.cells[1], 130, 120, 130, 120)
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        rr = p1.add_run(label.upper())
        rr.font.size = Pt(8.5)
        rr.font.bold = True
        rr.font.color.rgb = RGBColor.from_string(WHITE)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        rr = p2.add_run(value)
        rr.font.size = Pt(9.5)
        rr.font.bold = label == "Release truth" or row_idx == 0
        rr.font.color.rgb = RGBColor.from_string(WHITE if row_idx == 0 else INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Classification: project working document - public, representative and simulated data only")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def add_toc(doc: Document):
    p = doc.add_paragraph("Contents", style="Heading 1")
    set_keep(p, True)
    source_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    titles = [line[2:].strip() for line in source_lines if line.startswith("# ")][1:]
    for title in titles:
        entry = doc.add_paragraph()
        entry.paragraph_format.left_indent = Inches(0.12)
        entry.paragraph_format.first_line_indent = Inches(-0.12)
        entry.paragraph_format.space_after = Pt(2.5)
        run = entry.add_run(title)
        run.font.name = "Calibri"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()


def table_widths(rows):
    cols = len(rows[0])
    total = 9360
    max_lens = []
    for idx in range(cols):
        max_lens.append(max(8, min(52, max(len(re.sub(r"[`*]", "", row[idx])) for row in rows))))
    min_widths = []
    for idx in range(cols):
        max_word = max(
            len(word)
            for row in rows
            for word in re.sub(r"[`*]", "", row[idx]).replace("/", " / ").split()
        )
        min_widths.append(max(1050, min(2400, 320 + 80 * max_word)))
    raw = [max(min_widths[i], int(total * n / sum(max_lens))) for i, n in enumerate(max_lens)]
    while sum(raw) > total:
        candidates = [i for i in range(cols) if raw[i] > min_widths[i]]
        if not candidates:
            scale = total / sum(raw)
            raw = [max(900, int(width * scale)) for width in raw]
            break
        largest = max(candidates, key=lambda i: raw[i] - min_widths[i])
        raw[largest] -= min(50, sum(raw) - total, raw[largest] - min_widths[largest])
    if sum(raw) < total:
        raw[-1] += total - sum(raw)
    elif sum(raw) > total:
        raw[-1] -= sum(raw) - total
    return raw


def add_markdown_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    widths = table_widths(rows)
    set_table_width(table, widths)
    for r_idx, row_values in enumerate(rows):
        row = table.rows[r_idx]
        if r_idx == 0:
            repeat_table_header(row)
        # Body rows remain splittable because some reference cells can exceed
        # the usable page height in Word.
        for c_idx, value in enumerate(row_values):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, TABLE_HEADER)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value.strip(), color=NAVY if r_idx == 0 else INK, size=8.0 if len(rows[0]) >= 4 else 8.5)
            for run in p.runs:
                run.font.name = "Calibri"
                if r_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(WHITE)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cells = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        idx += 1
    return rows, idx


def add_code_block(doc, code_lines):
    for idx, line in enumerate(code_lines):
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.space_before = Pt(0 if idx else 5)
        p.paragraph_format.space_after = Pt(0 if idx < len(code_lines) - 1 else 6)
        set_cell_like_shading(p, PALE)
        p.add_run(line or " ")


def set_cell_like_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), GOLD)
    borders.append(left)
    p_pr.append(borders)


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, 120, 180, 120, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, color=INK, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_image(doc, alt, src):
    path = (SOURCE.parent / src).resolve()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(6.45))
    set_alt_text(shape, alt, alt + ". Diagram described in the surrounding section text.")
    cap = doc.add_paragraph(alt, style="Caption")
    set_keep(cap, False)


def parse_markdown(doc: Document, numbering):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Document control")
    lines = lines[start:]
    idx = 0
    current_decimal_num = numbering["decimal"]
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if not stripped or stripped == "---":
            idx += 1
            continue
        if stripped.startswith("```"):
            code_lines = []
            idx += 1
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code_lines.append(lines[idx])
                idx += 1
            add_code_block(doc, code_lines)
            idx += 1
            continue
        img_match = re.fullmatch(r"!\[([^\]]+)\]\(([^)]+)\)", stripped)
        if img_match:
            add_image(doc, img_match.group(1), img_match.group(2))
            idx += 1
            continue
        if stripped.startswith("|") and idx + 1 < len(lines) and re.match(r"^\|?\s*:?-{3,}", lines[idx + 1].strip()):
            rows, idx = parse_table(lines, idx)
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, title)
            set_keep(p, True)
            idx += 1
            continue
        if stripped.startswith("> "):
            add_note(doc, stripped[2:])
            idx += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph()
            set_num(p, numbering["bullet"])
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            add_inline(p, bullet.group(1))
            set_keep(p)
            idx += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            if int(numbered.group(1)) == 1:
                current_decimal_num = new_number_instance(doc, numbering["decimal"])
            p = doc.add_paragraph()
            set_num(p, current_decimal_num)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            add_inline(p, numbered.group(2))
            set_keep(p)
            idx += 1
            continue

        paragraph_lines = [stripped]
        idx += 1
        while idx < len(lines):
            candidate = lines[idx].strip()
            if not candidate or candidate == "---" or candidate.startswith(("#", "|", "```", "> ", "- ")) or re.match(r"^\d+\.\s+", candidate) or re.fullmatch(r"!\[[^\]]+\]\([^)]+\)", candidate):
                break
            paragraph_lines.append(candidate)
            idx += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_lines))
        set_keep(p)


def enable_update_fields(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def scrub_metadata(doc: Document):
    props = doc.core_properties
    props.title = "Corporate Wallet Digital Twin V3.1 - Complete System Dossier"
    props.subject = "Authoritative implementation snapshot as at 12 August 2026"
    props.author = "Corporate Wallet Digital Twin team - Christopher Koen"
    props.keywords = "corporate banking, wallet, evidence, Bayesian, economics, timing, GenAI, AWS, Databricks"
    props.comments = "Client demo READY; bank production NOT_PROMOTABLE."
    props.last_modified_by = "Corporate Wallet Digital Twin team - Christopher Koen"


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    build_architecture_diagram(ASSET_DIR / "v3_target_architecture.png")
    build_ladder_diagram(ASSET_DIR / "v3_evidence_claim_ladder.png")
    build_decision_architecture(ASSET_DIR / "v3_decision_architecture.png")

    doc = Document()
    configure_styles(doc)
    section = configure_section(doc)
    doc.settings.odd_and_even_pages_header_footer = True
    add_running_header_footer(section)
    numbering = add_numbering_definition(doc)
    add_cover(doc)
    add_toc(doc)
    parse_markdown(doc, numbering)
    scrub_metadata(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()

