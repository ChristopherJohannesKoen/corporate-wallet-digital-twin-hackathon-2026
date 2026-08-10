from __future__ import annotations

"""Build the literature-grounded V3 technical foundations white paper.

The body uses the exact compact_reference_guide preset and reuses the hardened
OOXML utilities from the system-dossier builder.  The cover is a named
editorial_cover override.
"""

import re
import sys
from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
import build_v3_system_dossier as base  # noqa: E402


SOURCE = ROOT / "docs" / "Corporate_Wallet_Digital_Twin_V3_Technical_Foundations.md"
OUTPUT = ROOT / "deliverables" / "Corporate_Wallet_Digital_Twin_V3_Technical_Foundations.docx"
ASSET_DIR = ROOT / "deliverables" / "assets"


def configure_styles(doc: Document) -> None:
    """Apply the exact compact_reference_guide typography tokens."""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(base.INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"] if "Caption" in styles else styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = "Calibri"
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(base.MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)

    code = styles["Code Block"] if "Code Block" in styles else styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(base.NAVY)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.0


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(40)
    r = p.add_run("STANDARD BANK HACKATHON 2026  /  TECHNICAL WHITE PAPER")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(base.GOLD)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Corporate Wallet\nDigital Twin V3")
    r.font.name = "Calibri"
    r.font.size = Pt(31)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(base.NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("Technical Foundations, Statistical Theory\nand Production Engineering")
    r.font.name = "Calibri"
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor.from_string(base.BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.62)
    p.paragraph_format.right_indent = Inches(0.62)
    p.paragraph_format.space_after = Pt(34)
    r = p.add_run(
        "A literature-grounded specification of the epistemic model, data contracts, "
        "wallet inference, latent-network reconstruction, PU learning, Bayesian temporal dynamics, "
        "robust decision optimization, value of information, controlled GenAI and production engineering."
    )
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(base.INK)

    table = doc.add_table(rows=7, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = [
        ("DOCUMENT CONTROL", "VALUE"),
        ("Team", "Corporate Wallet Digital Twin"),
        ("Team member", "Christopher Koen"),
        ("Version", "3.0.0"),
        ("Technical baseline", "Corporate Wallet Digital Twin V3"),
        ("As of", "10 August 2026"),
        ("Release truth", "V3 Decision Lab ready / bank production not promotable"),
    ]
    base.set_table_width(table, [2200, 7160])
    base.repeat_table_header(table.rows[0])
    for row_index, (row, (label, value)) in enumerate(zip(table.rows, values)):
        base.set_cell_shading(row.cells[0], base.NAVY)
        base.set_cell_shading(row.cells[1], base.NAVY if row_index == 0 else base.PALE)
        for cell in row.cells:
            base.set_cell_margins(cell, 130, 120, 130, 120)
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        rr = p1.add_run(label.upper())
        rr.font.size = Pt(8.5)
        rr.font.bold = True
        rr.font.color.rgb = RGBColor.from_string(base.WHITE)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        rr = p2.add_run(value)
        rr.font.size = Pt(9.5)
        rr.font.bold = row_index in (0, 6)
        rr.font.color.rgb = RGBColor.from_string(base.WHITE if row_index == 0 else base.INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Classification: technical design record - public, representative and simulated data only")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(base.MUTED)
    doc.add_page_break()


def add_running_header_footer(section) -> None:
    def populate_header(header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("CORPORATE WALLET DIGITAL TWIN V3  /  TECHNICAL FOUNDATIONS")
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(base.MUTED)

    def populate_footer(footer):
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("LITERATURE-GROUNDED DESIGN  •  10 AUGUST 2026     ")
        r.font.name = "Calibri"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(base.MUTED)
        base.add_field(p, "PAGE")

    populate_header(section.header)
    populate_header(section.even_page_header)
    populate_footer(section.footer)
    populate_footer(section.even_page_footer)


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph("Contents", style="Heading 1")
    base.set_keep(p, True)
    titles = [line[3:].strip() for line in SOURCE.read_text(encoding="utf-8").splitlines() if line.startswith("## ")]

    # A fixed two-column table is used instead of a field-based TOC.  It is
    # deterministic in headless LibreOffice/Word rendering and avoids the
    # occasional overlapping paragraphs produced by automatic TOC fields.
    midpoint = (len(titles) + 1) // 2
    columns = (titles[:midpoint], titles[midpoint:])
    table = doc.add_table(rows=midpoint + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    base.set_table_width(table, [4680, 4680])
    base.repeat_table_header(table.rows[0])

    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        borders.append(border)
    tbl_pr.append(borders)

    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            base.set_cell_margins(cell, 38, 120, 38, 120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.keep_together = True
            if row_index == 0:
                label = ("Foundations and analytical design", "Architecture, controls and appendices")[column_index]
                run = p.add_run(label.upper())
                run.font.name = "Calibri"
                run.font.size = Pt(7.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(base.MUTED)
                continue
            title_index = row_index - 1
            if title_index < len(columns[column_index]):
                run = p.add_run(columns[column_index][title_index])
                run.font.name = "Calibri"
                run.font.size = Pt(9.2)
                run.font.color.rgb = RGBColor.from_string(base.NAVY)
    doc.add_page_break()


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    base.set_table_width(table, [9360])
    base.repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    base.set_cell_shading(cell, base.CALLOUT_FILL)
    base.set_cell_margins(cell, 120, 120, 120, 120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    base.add_inline(p, text, color=base.INK, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_markdown(doc: Document, numbering) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.strip() == "## Document control")
    lines = lines[start:]
    index = 0
    current_decimal = numbering["decimal"]
    while index < len(lines):
        stripped = lines[index].strip()
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            base.add_code_block(doc, code_lines)
            index += 1
            continue
        image_match = re.fullmatch(r"!\[([^\]]+)\]\(([^)]+)\)", stripped)
        if image_match:
            base.add_image(doc, image_match.group(1), image_match.group(2))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-{3,}", lines[index + 1].strip()):
            rows, index = base.parse_table(lines, index)
            base.add_markdown_table(doc, rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            p = doc.add_paragraph(style=f"Heading {level}")
            base.add_inline(p, heading.group(2))
            base.set_keep(p, True)
            index += 1
            continue
        if stripped.startswith("> "):
            base.add_note(doc, stripped[2:])
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph()
            base.set_num(p, numbering["bullet"])
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            base.add_inline(p, bullet.group(1))
            base.set_keep(p)
            index += 1
            continue
        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            if int(numbered.group(1)) == 1:
                current_decimal = base.new_number_instance(doc, numbering["decimal"])
            p = doc.add_paragraph()
            base.set_num(p, current_decimal)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.25
            base.add_inline(p, numbered.group(2))
            base.set_keep(p)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if (
                not candidate
                or candidate == "---"
                or candidate.startswith(("#", "|", "```", "> ", "- "))
                or re.match(r"^\d+\.\s+", candidate)
                or re.fullmatch(r"!\[[^\]]+\]\([^)]+\)", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        p = doc.add_paragraph()
        base.add_inline(p, " ".join(paragraph_lines))
        base.set_keep(p)


def build_inference_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 1040), base.pil_color(base.WHITE))
    draw = ImageDraw.Draw(img)
    draw.text((80, 45), "Wallet inference as a governed measurement system", font=base.font(40, True), fill=base.pil_color(base.NAVY))
    draw.text((80, 98), "Direct observations constrain the latent wallet; proxies inform it with tier-dependent error", font=base.font(24), fill=base.pil_color(base.MUTED))
    boxes = [
        ((70, 230, 390, 465), "Bank observation", ["Product activity y", "Balances and fees", "Event time / as-of"]),
        ((70, 590, 390, 825), "External evidence", ["E1 public anchors", "E2 attestations", "E3 multibank share"]),
        ((545, 335, 920, 720), "Latent state", ["Total wallet W", "Bank share s = y/W", "Competitor residual 1-s"]),
        ((1065, 230, 1435, 465), "Identified set", ["Deterministic [L,U]", "No distribution assumed", "Assumption-light"]),
        ((1065, 590, 1435, 825), "Posterior", ["Product Beta model", "Anchor likelihood/pooling", "Reproducible draws"]),
        ((1580, 335, 1840, 720), "Decision layer", ["Scenario frontier", "Timing probability", "Eligibility gates"]),
    ]
    fills = [base.SKY, "FFF5DA", "E7F4EE", base.PALE, "F3EAF8", "FCEAEA"]
    for (box, title, lines), fill in zip(boxes, fills):
        base.rounded_box(draw, box, fill, base.NAVY)
        base.centered_lines(draw, box, [title] + lines)
    for start, end in [((390, 350), (545, 450)), ((390, 705), (545, 610)), ((920, 460), (1065, 350)), ((920, 600), (1065, 705)), ((1435, 350), (1580, 455)), ((1435, 705), (1580, 600))]:
        base.arrow(draw, start, end)
    draw.text((80, 930), "Rule: bounds, posterior estimates and commercial scenarios remain separately labelled and separately validated.", font=base.font(23, True), fill=base.pil_color(base.NAVY))
    img.save(path, quality=95)


def build_architecture_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 1040), base.pil_color(base.WHITE))
    draw = ImageDraw.Draw(img)
    draw.text((80, 48), "Corporate Wallet Digital Twin V3 — target production flow", font=base.font(40, True), fill=base.pil_color(base.NAVY))
    draw.text((80, 102), "Private, point-in-time and governed from source to outcome", font=base.font(24), fill=base.pil_color(base.MUTED))
    boxes = {
        "sources": (70, 210, 360, 780), "ingest": (440, 210, 735, 455), "lake": (440, 535, 735, 780),
        "models": (820, 210, 1140, 455), "evidence": (820, 535, 1140, 780), "bff": (1225, 210, 1535, 455),
        "events": (1225, 535, 1535, 780), "controls": (1615, 210, 1850, 780),
    }
    fills = {"sources": base.PALE, "ingest": base.SKY, "lake": base.SKY, "models": "E7F4EE", "evidence": "FFF5DA", "bff": "E7F4EE", "events": "F3EAF8", "controls": "FCEAEA"}
    text_map = {
        "sources": ["Sources", "Bank activity & balances", "CRM and finance", "Public documents", "Consented multibank data"],
        "ingest": ["AWS ingestion", "Contract validation", "Quarantine", "Immutable source hashes"],
        "lake": ["S3 + Delta Lake", "Raw → conformed → curated", "Point-in-time features", "Object-locked snapshots"],
        "models": ["Decision services", "Bounds + posterior + shadow", "PU + temporal dynamics", "CVaR portfolio + VOI"],
        "evidence": ["Evidence + GenAI", "Extraction candidates", "Four-eyes approval", "Controlled narration"],
        "bff": ["Entitled workbench", "ABAC server-side APIs", "Separate claim layers", "CRM adapter"],
        "events": ["MSK learning loop", "Eligibility and assignment", "Interaction and action", "Outcome and access events"],
        "controls": ["Control plane", "SSO + OPA + Unity Catalog", "KMS + MLflow registry", "OpenTelemetry + SIEM", "Monitoring + rollback"],
    }
    for name, box in boxes.items():
        base.rounded_box(draw, box, fills[name], base.NAVY if name in {"sources", "controls"} else base.BLUE)
        base.centered_lines(draw, box, text_map[name])
    for start, end in [((360, 335), (440, 335)), ((587, 455), (587, 535)), ((735, 335), (820, 335)), ((735, 655), (820, 655)), ((1140, 335), (1225, 335)), ((1380, 455), (1380, 535)), ((1535, 655), (1615, 655))]:
        base.arrow(draw, start, end)
    base.arrow(draw, (1615, 335), (1535, 335), color=base.GOLD)
    draw.text((70, 900), "Release truth", font=base.font(25, True), fill=base.pil_color(base.NAVY))
    draw.text((235, 900), "Client demo READY", font=base.font(25, True), fill=base.pil_color(base.GREEN))
    draw.text((515, 900), "Bank production NOT PROMOTABLE until authority-dependent gates pass", font=base.font(25, True), fill=base.pil_color(base.RED))
    img.save(path, quality=95)


def build_ladder_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 900), base.pil_color(base.WHITE))
    draw = ImageDraw.Draw(img)
    draw.text((80, 45), "Evidence and claim ladder", font=base.font(40, True), fill=base.pil_color(base.NAVY))
    draw.text((80, 98), "Stronger evidence permits stronger labels — never automatic promotion", font=base.font(24), fill=base.pil_color(base.MUTED))
    tiers = [("E0", "Governed prior", "Prior-led posterior or scenario"), ("E1", "Audited public evidence", "Noisy or censored anchor"), ("E2", "Client / RM attestation", "Client-validated inference"), ("E3", "Multibank observation", "Measured share eligible"), ("E4", "Reconciled outcome", "Reconciled or causal evidence")]
    colors = ["EEF2F5", "DFECF7", "D7EBDD", "F8E9C5", "F1D6D6"]
    for i, ((tier, name, claim), fill) in enumerate(zip(tiers, colors)):
        x1, y1 = 80 + i * 345, 600 - i * 100
        box = (x1, y1, x1 + 330, 780)
        base.rounded_box(draw, box, fill, base.NAVY, radius=18, width=3)
        draw.text((x1 + 20, y1 + 18), tier, font=base.font(34, True), fill=base.pil_color(base.NAVY))
        draw.text((x1 + 20, y1 + 65), name, font=base.font(22, True), fill=base.pil_color(base.BLUE))
        draw.multiline_text((x1 + 20, y1 + 105), base.wrap_text(draw, claim, base.font(19), 290), font=base.font(19), fill=base.pil_color(base.INK), spacing=6)
    draw.text((82, 828), "Observed → identified bound → posterior → scenario → causal", font=base.font(23, True), fill=base.pil_color(base.NAVY))
    img.save(path, quality=95)


def build_learning_diagram(path: Path) -> None:
    img = Image.new("RGB", (1900, 940), base.pil_color(base.WHITE))
    draw = ImageDraw.Draw(img)
    draw.text((80, 45), "From eligibility to causal evidence", font=base.font(40, True), fill=base.pil_color(base.NAVY))
    draw.text((80, 98), "The event stream preserves the counterfactual denominator before any recommendation is displayed", font=base.font(24), fill=base.pil_color(base.MUTED))
    labels = [
        ("Eligible", "Every candidate logged"),
        ("Assigned", "Cluster encouragement"),
        ("Displayed", "Exposure is observed"),
        ("Action", "Qualified RM action"),
        ("Outcome", "Milestone / economics"),
    ]
    xs = [70, 430, 790, 1150, 1510]
    colors = [base.SKY, "FFF5DA", "E7F4EE", "F3EAF8", "FCEAEA"]
    for x, (title, body), fill in zip(xs, labels, colors):
        box = (x, 280, x + 310, 510)
        base.rounded_box(draw, box, fill, base.NAVY)
        base.centered_lines(draw, box, [title, body])
    for x in xs[:-1]:
        base.arrow(draw, (x + 310, 395), (x + 360, 395))
    base.rounded_box(draw, (240, 650, 1660, 840), base.PALE, base.BLUE)
    base.centered_lines(draw, (240, 650, 1660, 840), ["Estimands and gates", "Primary: intention-to-treat risk difference", "Secondary: Wald complier effect only with a valid first stage", "Heterogeneity and policy learning only after overlap, sample-size and independent-validation gates"])
    base.arrow(draw, (945, 510), (945, 650), color=base.GOLD)
    img.save(path, quality=95)


def scrub_metadata(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Corporate Wallet Digital Twin V3 — Technical Foundations, Statistical Theory and Production Engineering"
    props.subject = "Literature-grounded V3 technical design specification as at 10 August 2026"
    props.author = "Corporate Wallet Digital Twin team - Christopher Koen"
    props.keywords = "corporate banking, wallet, Bayesian measurement, partial identification, causal inference, GenAI, AWS, Databricks"
    props.comments = "Client demonstration ready; bank production not promotable."
    props.last_modified_by = "Corporate Wallet Digital Twin team - Christopher Koen"


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    build_inference_diagram(ASSET_DIR / "v3_wallet_inference.png")
    build_learning_diagram(ASSET_DIR / "v3_causal_learning.png")

    base.SOURCE = SOURCE
    base.OUTPUT = OUTPUT
    base.OUT_DIR = OUTPUT.parent
    base.ASSET_DIR = ASSET_DIR
    base.configure_styles = configure_styles
    base.add_cover = add_cover
    base.add_toc = add_toc
    base.add_note = add_note
    base.add_running_header_footer = add_running_header_footer
    base.parse_markdown = parse_markdown
    base.scrub_metadata = scrub_metadata
    base.build_architecture_diagram = build_architecture_diagram
    base.build_ladder_diagram = build_ladder_diagram
    base.main()


if __name__ == "__main__":
    main()
